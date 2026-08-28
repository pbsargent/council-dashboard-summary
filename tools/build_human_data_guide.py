#!/usr/bin/env python3
from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs"
OUT_FILE = OUT_DIR / "Council-Dashboard-Summary-Source-and-Calculation-Guide.docx"
LATEST_JSON = ROOT / "data" / "latest.json"
MONDAY_JSON = ROOT / "data" / "monday-latest.json"
LOGO = ROOT / "assets" / "cac-logo.png"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(18, 32, 51)
MUTED = RGBColor(83, 97, 116)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
LINE = "CAD6E2"


def load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def fmt_num(value) -> str:
    if value is None:
        return "n/a"
    try:
        return f"{float(value):,.0f}"
    except (TypeError, ValueError):
        return str(value)


def fmt_pct(value) -> str:
    if value is None:
        return "n/a"
    try:
        return f"{float(value) * 100:.1f}%"
    except (TypeError, ValueError):
        return str(value)


def weighted_section_rate(sections: list[dict], section_name: str, field: str):
    section = next((item for item in sections if item.get("section") == section_name), None)
    if not section:
        return None
    measured_rows = [
        row
        for row in section.get("rows", [])
        if row.get(field) is not None and row.get("units")
    ]
    total_units = sum(float(row["units"]) for row in measured_rows)
    if not total_units:
        return None
    return sum(float(row[field]) * float(row["units"]) for row in measured_rows) / total_units


def short_date(value: str | None) -> str:
    if not value:
        return "n/a"
    try:
        return datetime.fromisoformat(value.replace("Z", "+00:00")).strftime("%b %-d, %Y")
    except ValueError:
        return value


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color=LINE, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, col_widths_dxa: list[int], indent_dxa=120):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(col_widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in col_widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, col_widths_dxa[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_row_pagination(row, repeat_header=False):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
    if repeat_header:
        tbl_header = tr_pr.find(qn("w:tblHeader"))
        if tbl_header is None:
            tbl_header = OxmlElement("w:tblHeader")
            tr_pr.append(tbl_header)
        tbl_header.set(qn("w:val"), "true")


def paragraph_border_bottom(paragraph, color="2E74B5", size="12", space="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)


def add_heading(doc: Document, text: str, level: int = 1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    run = paragraph.add_run(text)
    return paragraph


def add_body(doc: Document, text: str):
    paragraph = doc.add_paragraph(style="Body Text")
    paragraph.add_run(text)
    return paragraph


def add_bullet(doc: Document, text: str):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.add_run(text)
    return paragraph


def add_reference_bullet(doc: Document, text: str):
    paragraph = doc.add_paragraph(style="Body Text")
    paragraph.paragraph_format.left_indent = Inches(0.28)
    paragraph.paragraph_format.first_line_indent = Inches(-0.18)
    paragraph.add_run("•  ")
    paragraph.add_run(text)
    return paragraph


def add_step(doc: Document, text: str):
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.add_run(text)
    return paragraph


def add_callout(doc: Document, title: str, body: str):
    table = doc.add_table(rows=1, cols=1)
    set_row_pagination(table.rows[0])
    set_table_width(table, [9360], indent_dxa=120)
    set_table_borders(table, color="D8E0EA", size="4")
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    title_run = p.add_run(title)
    set_run_font(title_run, size=10.5, color=DARK_BLUE, bold=True)
    body_p = cell.add_paragraph()
    body_p.paragraph_format.space_after = Pt(0)
    body_run = body_p.add_run(body)
    set_run_font(body_run, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_width(table, widths, indent_dxa=120)
    set_table_borders(table)
    header_cells = table.rows[0].cells
    set_row_pagination(table.rows[0], repeat_header=True)
    for idx, header in enumerate(headers):
        set_cell_shading(header_cells[idx], LIGHT_BLUE)
        paragraph = header_cells[idx].paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(header)
        set_run_font(run, size=9.5, color=INK, bold=True)
    for row in rows:
        table_row = table.add_row()
        set_row_pagination(table_row)
        cells = table_row.cells
        for idx, value in enumerate(row):
            paragraph = cells[idx].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(str(value))
            set_run_font(run, size=9.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def setup_document(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    body = styles["Body Text"]
    body.base_style = normal
    body.paragraph_format.space_after = Pt(6)
    body.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def add_running_furniture(doc: Document):
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.text = ""
    left = p.add_run("Council Dashboard Summary")
    set_run_font(left, size=9, color=MUTED, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph_border_bottom(p, color="D8E0EA", size="4", space="4")

    footer = section.footer
    p = footer.paragraphs[0]
    p.text = ""
    run = p.add_run("Source and Calculation Guide")
    set_run_font(run, size=9, color=MUTED)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def add_cover(doc: Document, latest: dict, monday: dict):
    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        logo = p.add_run().add_picture(str(LOGO), width=Inches(2.2))
        logo._inline.docPr.set("title", "Capitol Area Council")
        logo._inline.docPr.set("descr", "Capitol Area Council logo")
        p.paragraph_format.space_after = Pt(22)

    kicker = doc.add_paragraph()
    run = kicker.add_run("Council Dashboard Summary")
    set_run_font(run, size=10, color=BLUE, bold=True)
    kicker.paragraph_format.space_after = Pt(5)

    title = doc.add_paragraph()
    title_run = title.add_run("Source and Calculation Guide")
    set_run_font(title_run, size=25, color=INK, bold=True)
    title.paragraph_format.space_after = Pt(4)

    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run("A plain-English guide to where dashboard values come from, how they are refreshed, how major metrics are calculated, and how the in-dashboard help works.")
    set_run_font(subtitle_run, size=12.5, color=MUTED)
    subtitle.paragraph_format.space_after = Pt(18)
    paragraph_border_bottom(subtitle, color="2E74B5", size="10", space="8")

    dashboard = latest["dashboard"]
    cst = latest["cst"]
    add_table(
        doc,
        ["Item", "Current source / value"],
        [
            ["Data current", short_date(latest.get("generated_at"))],
            ["Dashboard workbook", dashboard.get("source_name", "n/a")],
            ["CST workbook", cst.get("source_name", "n/a")],
            ["monday.com workbook", monday.get("source_workbook") or monday.get("generated_from", "n/a")],
            ["Guide generated", short_date(datetime.now().isoformat())],
            ["Published site", "https://pbsargent.github.io/council-dashboard-summary/"],
        ],
        [2400, 6960],
    )

    add_callout(
        doc,
        "Reader promise",
        "This guide avoids code-level detail unless a formula matters. The dashboard's ? buttons provide brief panel-level help; this document provides the deeper source and calculation context. For exact implementation and edge cases, use the Markdown data dictionary in the repository.",
    )

    doc.add_page_break()


def build_doc():
    latest = load_json(LATEST_JSON)
    monday = load_json(MONDAY_JSON)
    council = latest["dashboard"]["council"]
    boards = monday["boards"]
    council_retention = weighted_section_rate(
        latest["dashboard"].get("unit_metric_compare", []),
        "All Units",
        "retention_rate",
    )

    doc = Document()
    setup_document(doc)
    add_running_furniture(doc)
    add_cover(doc, latest, monday)

    add_heading(doc, "1. What The Dashboard Is", 1)
    add_body(
        doc,
        "The Council Dashboard Summary is a static web dashboard for Capitol Area Council operating review. It brings together membership, retention, unit health, renewal, training, camping readiness, safeguarding, commissioner coverage, Service Area context, CST comparison, recruitment, and monday.com operating context.",
    )
    add_body(
        doc,
        "The website itself is static. It does not connect directly to Google Drive or monday.com when someone opens it. Instead, a daily local refresh process creates compact JSON and JavaScript data bundles in an isolated staging tree, verifies the payload, and deploys it through GitHub Pages Actions. Generated daily data is not committed to the source repository.",
    )
    add_body(
        doc,
        "Most major panels include a circular ? help control. Hovering, focusing, or clicking that control opens a short explanation of the panel's source and meaning. Those panel notes are intentionally brief; this guide remains the fuller reference.",
    )

    add_heading(doc, "At-a-glance data currently published", 2)
    add_table(
        doc,
        ["Area", "Current value"],
        [
            ["Youth membership", fmt_num(council.get("members"))],
            ["Volunteers", fmt_num(council.get("volunteers"))],
            ["Retention", fmt_pct(council_retention)],
            ["Units", fmt_num(council.get("units"))],
            ["At-risk units", f"{fmt_num(council.get('at_risk_units'))} ({fmt_pct(council.get('at_risk_rate'))})"],
            ["Healthy units", f"{fmt_num(council.get('healthy_units'))} ({fmt_pct(council.get('healthy_rate'))})"],
            ["Commissioners", fmt_num(council.get("commissioners"))],
            ["Unit commissioners", fmt_num(council.get("unit_commissioners"))],
            ["Service Areas", fmt_num(len(latest["dashboard"].get("service_areas", [])))],
            ["Training rows", fmt_num(len(latest["dashboard"].get("training_people", [])))],
            ["Hot prospects", fmt_num(boards["prospects"].get("items"))],
            ["Renewal rows", fmt_num(boards["renewals"].get("items"))],
            ["School rows", fmt_num(boards["schools"].get("items"))],
            ["Popcorn units", fmt_num(boards.get("popcorn", {}).get("items"))],
            [
                "Popcorn committed",
                f"{fmt_num(boards.get('popcorn', {}).get('committed'))} "
                f"({fmt_pct(boards.get('popcorn', {}).get('participation_rate'))})",
            ],
        ],
        [3000, 6360],
    )

    add_heading(doc, "2. Where The Data Comes From", 1)
    add_body(doc, "The dashboard has six main source lanes.")
    add_table(
        doc,
        ["Source lane", "Used for", "Published as"],
        [
            ["Council dashboard workbook", "Membership, unit health, district training rates, commissioner objectives, person-level training, and commissioner roster.", "`data/latest.json`"],
            ["CST comparison workbook", "Service Territory comparison metrics and Capitol Area Council comparison fields.", "`data/latest.json`"],
            ["monday.com daily workbook", "Hot prospects, 2026 renewal status, schools, Total Available Youth, and Popcorn commitments/readiness.", "`data/monday-latest.json`"],
            ["Unit Level Metrics workbook", "Unit-level youth, health, growth, training, SYT, and commissioner-assignment detail by program.", "`data/unit-level-latest.*`"],
            ["Cub Scout JSN board", "School recruiting plans, scheduled recruitments, location and date patterns, and recruiting-material totals.", "`data/fall-recruitment-latest.js`"],
            ["Service Areas board", "Authoritative Service Area hierarchy, Field Directors, district professionals, Volunteer Chairs, and District Commissioners.", "`data/latest.json`"],
        ],
        [2200, 5000, 2160],
    )
    add_body(
        doc,
        "The monday.com daily workbook is preferred because it contains richer detail rows. If that workbook is not available, the automation can fall back to the monday.com API, but the API fallback is less detailed than the workbook export. The Council Dashboard Reports drive also contains Unit Level Metrics captures; those are used only for Unit Level detail and must not be selected as the main Dashboard - CAC workbook.",
    )

    add_heading(doc, "3. How The Daily Refresh Works", 1)
    add_body(
        doc,
        "The consolidated CAC Dashboard Platform is the sole production writer. One scheduled macOS LaunchAgent runs the platform orchestrator, which refreshes source data and publishes the Council Summary and Commissioner sites from their active working copies. The Commissioner Dashboard remains a separate portal that reads the same canonical Council Summary data.",
    )
    for step in [
        "Capture the current monday.com Service Areas hierarchy and district leadership fields for this run.",
        "Find the newest `*_Dashboard - CAC.xlsx` Council dashboard workbook in the Council Dashboard Reports shared drive.",
        "Ignore `*_CAC - Unit Metric Scorecard.xlsx` for the CAC dashboard source selection, even if that Unit Level Metrics workbook is the newest workbook in the same shared drive.",
        "Find the newest CST7 workbook in the Council Metric Reports shared drive.",
        "Build `data/latest.json`, including Service Area rollups and a dated working archive outside the Pages source tree.",
        "Find the newest monday.com export workbook in the Council monday.com Reports shared drive.",
        "Build `data/monday-latest.json`, falling back to the monday.com API only if needed.",
        "Rebuild Unit-Level Detail, Renewal Status, Units-Youth history, and the Cub Scout JSN data bundle.",
        "Run the site-structure validator and stage the complete static site in an isolated Pages tree.",
        "Deploy the Council Summary site tree as a checksum-verified GitHub Pages artifact without committing generated data.",
        "Deploy the Commissioner Dashboard as a verified Pages artifact only when its static payload changed.",
        "Verify the GitHub Pages deployment and report the refresh result.",
    ]:
        add_step(doc, step)
    add_callout(
        doc,
        "What changes daily",
        "The daily automation updates data files. It does not rewrite dashboard code or formulas. A formula or page layout changes only when the website code is edited and committed.",
    )
    add_callout(
        doc,
        "Single-writer publishing",
        "The consolidated CAC Dashboard Platform is the only production deployer. Source changes use ordinary linear commits; daily generated data uses verified Pages artifacts without force pushes.",
    )
    add_callout(
        doc,
        "How freshness times display",
        "The lower-left dashboard timestamp is shown in the viewer's local timezone. Council dashboard timestamps that do not include a timezone are interpreted as America/Chicago source time first; monday.com timestamps that include UTC or another offset are treated as exact instants and then displayed locally for the viewer.",
    )
    add_callout(
        doc,
        "Automation paths",
        "The active platform root is /Users/petersargent/CACDashboardPlatform. The sole LaunchAgent is /Users/petersargent/Library/LaunchAgents/com.cac.dashboard.macpro-daily.plist. The platform orchestrator is tools/daily_build_publish.zsh; the Council Summary updater is sites/council-dashboard-summary/update_daily.zsh.",
    )

    add_heading(doc, "4. Build And Data Acquisition Requirements", 1)
    add_body(
        doc,
        "A technical maintainer can rebuild a similar dashboard if they have the source workbooks, monday.com read access, a Python environment with the workbook/document libraries, a GitHub Pages repository, and an automation path that can build and deploy a verified static payload.",
    )
    add_table(
        doc,
        ["Requirement", "What must be available"],
        [
            ["Runtime", "Python with openpyxl for Excel parsing and python-docx for guide generation."],
            ["Google Drive access", "Shared drives named Council Dashboard Reports, Council Metric Reports, and Council monday.com Reports. Council Dashboard Reports contains both the Dashboard - CAC workbook and the Unit Level Metrics workbook."],
            ["Workbook patterns", "Use `*_Dashboard - CAC.xlsx` only for the council dashboard source, `*_CAC - Unit Metric Scorecard.xlsx` only for Unit Level detail, `*_CST7.xlsx` for CST metrics, and `*monday-export.xlsx` for monday.com detail."],
            ["monday.com access", "Daily export workbook preferred; API token fallback requires read access to the configured boards."],
            ["Publishing", "GitHub Pages repository on main branch with static HTML, CSS, JavaScript, assets, and JSON data."],
            ["Automation", "One LaunchAgent that refreshes source data and deploys generated data as verified Pages artifacts. Ordinary linear Git commits are reserved for source changes."],
            ["Service Area hierarchy", "Read access to monday.com board 18420160563 for Service Areas and district leadership."],
            ["Panel help", "Shared `panel-help.js` and dashboard CSS provide active hover, focus, and click/tap help popovers for the ? controls."],
        ],
        [2400, 6960],
    )
    add_callout(
        doc,
        "Technical runbook",
        "The repository file IMPLEMENTATION_RUNBOOK.md documents setup paths, source acquisition, manual refresh commands, schedule/log checks, validation steps, common failure modes, and guidance for rebuilding a similar dashboard.",
    )

    add_heading(doc, "5. How The Main Page Metrics Are Calculated", 1)
    add_table(
        doc,
        ["Displayed metric", "Plain-English calculation"],
        [
            ["Youth", "Sum of district youth membership from the Membership tab."],
            ["Retention", "Weighted average of district retention values from the selected Unit Metric Compare section, using units with a reported retention value as weights. The Council view uses All Units."],
            ["Units", "Sum of district unit counts from the Membership tab."],
            ["Average Metric", "District average metrics weighted by the number of units in each district."],
            ["Assigned", "Assigned units divided by total units. Assigned units come from the Assigned tab."],
            ["Training", "District all-scouter training rates weighted by units."],
            ["Volunteers", "Councilwide value published on the Membership tab, calculated by the source workbook as the count of unique nonblank Member IDs in the Training roster. It is not district- or program-filterable."],
            ["Youth / TAY", "Council: council youth divided by raw school-row TAY. Program view: actual program youth divided by estimated grade/age-eligible school TAY."],
            ["Popcorn Participation", "Units marked Committed divided by every unit row in the published Popcorn snapshot."],
        ],
        [2300, 7060],
    )

    add_heading(doc, "Master program filter", 2)
    add_body(
        doc,
        "The masthead filter carries Council, Packs, Troops, Crews, Ships, or Posts across linked pages. It uses direct unit-type fields where available and standardized unit names where a source does not publish a separate type field.",
    )
    add_table(
        doc,
        ["Source or panel", "Program-view behavior"],
        [
            ["Unit Level Metrics", "Uses the direct unit_type field to rebuild youth, units, health, growth, training, SYT, and assignment rollups."],
            ["Training and SYT", "Uses each published person's direct unit_type field."],
            ["Prospects", "Uses Potential Unit Type(s); a multi-type prospect can appear in more than one program view."],
            ["Renewal and Popcorn", "Parses program from standardized unit names; Posts are excluded from the published Popcorn population."],
            ["Units-Youth and CST", "Remain council-only except for source-native Pack/Troop connection fields."],
        ],
        [2600, 6760],
    )

    add_heading(doc, "District status", 2)
    add_body(doc, "Each district is labeled Needs Attention, Monitor, or On Track using threshold rules.")
    add_table(
        doc,
        ["Status", "Rule"],
        [
            ["Needs Attention", "YoY membership is below -10%, or SYT is below 80%, or at-risk unit rate is 55% or higher."],
            ["Monitor", "Training is below 65%, or at-risk unit rate is 40% or higher, or SYT is below 85%."],
            ["On Track", "None of the Needs Attention or Monitor rules apply."],
        ],
        [2100, 7260],
    )

    add_heading(doc, "6. What Each Dashboard Page Adds", 1)
    add_table(
        doc,
        ["Page", "What it answers"],
        [
            ["Overview", "What the Council's current KPI picture is and which operating areas need attention first."],
            ["Council Comparison", "How Capitol Area Council compares with other Service Territory 07 councils."],
            ["District Performance", "Which districts lead or lag across membership, unit health, training, SYT, and commissioner coverage."],
            ["Membership & Growth", "Where membership opportunity, TAY penetration, unit health risk, prospects, and renewals combine into priority signals."],
            ["Unit Health & Renewal", "Where unit health, assignment, and renewal follow-up require action."],
            ["People & Readiness", "Where leader training, safeguarding, and camping-readiness gaps are concentrated."],
            ["Training", "Which people are trained, which leaders are direct-contact, and where direct-contact training gaps exist."],
            ["SYT", "Whether direct-contact leaders have current SYT, Hazardous Weather, BALOO, and IOLS-related requirements."],
            ["Pack Camping Readiness", "Which Packs lack roster-level BALOO or current Hazardous Weather coverage and therefore need follow-up before overnight camping."],
            ["Troop Camping Readiness", "Which Troops lack the published IOLS signal or current Hazardous Weather coverage and therefore need follow-up before overnight camping."],
            ["Recruitment Pipeline", "Where prospect, renewal, and school operating follow-up is concentrated."],
            ["Cub Scout JSN", "How school recruiting plans, dates, locations, materials, and uncovered schools compare with the monday.com source dashboard."],
            ["Popcorn", "How unit commitments, goals, prior sales, onboarding, and training roll up from Service Area to District to Unit."],
            ["Unit Metrics", "How districts and unit sections compare across unit health, training, outdoor, advancement, and the current workbook retention metric."],
            ["Unit-Level Detail", "Which individual units and members drive program-specific youth, growth, training, SYT, health, and assignment results."],
            ["Renewal Status", "Which units are initiated, submitted, pending acceptance, posted, or otherwise need renewal follow-up."],
            ["Data & Help", "Which source workbooks are current, how values are calculated, and where to report a problem."],
        ],
        [2100, 7260],
    )
    add_body(
        doc,
        "Detail pages that compare district records also expose Service Area filters. District filters remain available inside the selected Service Area, and official district views exclude operational labels that are not part of the 12-district Council dashboard structure.",
    )

    add_heading(doc, "Cub Scout JSN interpretation", 2)
    add_body(
        doc,
        "The Cub Scout JSN page is rebuilt directly from monday.com board 18420720719. Its charts preserve the source dashboard's widget-specific filters rather than applying one page-wide population to every graphic.",
    )
    add_bullet(doc, "District recruitment includes dated recruitments and a No Date category.")
    add_bullet(doc, "School-planning gaps remain pie charts with their documented Cub-target, unit-association, plan-status, and district scopes.")
    add_bullet(doc, "Materials gauges total the full board and retain the source dashboard's fixed gauge maxima.")
    add_bullet(doc, "The former expenses gauge is not part of the current monday.com dashboard and is not published.")

    add_heading(doc, "Retention metric", 2)
    add_table(
        doc,
        ["Question", "Current dashboard definition"],
        [
            ["What is the formula?", "(Current members - members new in the prior 12 months) divided by same-month prior-year members."],
            ["What does it measure?", "The share of the prior-year membership base represented by current members after removing people who joined within the prior 12 months."],
            ["Can retention exceed 100%?", "Yes. Values above 100% are valid under this source-workbook formula and must not be capped."],
            ["How is it published?", "The dashboard preserves the workbook's decimal ratio for calculations and sorting, then rounds displayed retention to the nearest whole percent. For example, 117.5% displays as 118%."],
        ],
        [2600, 6760],
    )
    add_heading(doc, "7. Service Area And District Filtering", 1)
    add_body(
        doc,
        "Service Area is not inferred from workbook formulas. The daily build captures the authoritative hierarchy from monday.com board 18420160563. That board also supplies Field Directors, district professionals, Volunteer Chairs, and District Commissioners; its leadership values replace workbook values, including intentional blanks. The builder attaches the hierarchy to district rows, priority units, training people, commissioner records, Unit Metric Compare rows, renewal data, and monday.com contexts where an official district can be identified.",
    )
    add_table(
        doc,
        ["Service Area", "Field Director", "Official districts"],
        [
            ["Northern", "Justin Brundin", "Bee Cave; Chisholm Trail; Hill Country; North Shore"],
            ["Central", "Vicki Rosengarten", "Armadillo; Colorado River; Exploring; San Gabriel; Thunderbird"],
            ["Southern", "Ed Grune", "Live Oak; Sacred Springs; Waterloo"],
        ],
        [1700, 2200, 5460],
    )
    add_callout(
        doc,
        "How monday.com rows are assigned",
        "monday.com rows can contain multiple district labels. A row is attached to a Service Area when any listed official district belongs to that Service Area. Rows with only non-official or blank district labels remain outside official district and Service Area rollups.",
    )

    add_heading(doc, "Popcorn commitments hierarchy", 2)
    add_body(
        doc,
        "The Popcorn page uses one drill-down hierarchy: Service Area > District > Unit. Service Areas and Districts start collapsed. Expanding a Service Area reveals its District rollups; expanding a District reveals the corresponding unit follow-up rows inline.",
    )
    add_table(
        doc,
        ["Popcorn measure", "How it is calculated or displayed"],
        [
            ["Participation", "Committed unit rows divided by all unit rows in the selected population."],
            ["Committed goal", "Sum of 2026 goal values for committed units only."],
            ["Prior sales", "Sum of 2025 sales values for committed units only."],
            ["Goal vs 2025", "Committed goal compared with prior sales for committed units."],
            ["Onboarded", "Committed units whose onboarding progress is 11/11."],
            ["Unit trained", "Committed units marked trained in the Popcorn source."],
            ["Unit follow-up", "Commitment, goal, sales, onboarding, training, kernel, kickoff, and last-update context shown beneath the District."],
        ],
        [2600, 6760],
    )
    add_callout(
        doc,
        "Public-data privacy",
        "The published Popcorn rows intentionally exclude primary contact names, email addresses, and phone numbers. The normal scheduled source is the Popcorn Committments sheet in the daily monday.com workbook; the monday.com API is retained as a fallback.",
    )

    add_heading(doc, "8. TAY And Membership Opportunity", 1)
    add_body(
        doc,
        "TAY means Total Available Youth. It comes from the monday.com Schools export, not from the base Council dashboard workbook.",
    )
    add_bullet(doc, "Council Youth / TAY uses raw school-row TAY once per school row.")
    add_bullet(doc, "District Youth / TAY attributes a school's full TAY to each official Scouting District listed for that school.")
    add_bullet(doc, "Program Youth / TAY uses actual youth from units of the selected type and estimates eligible TAY by allocating each school's total evenly across its published grade or age span.")
    add_bullet(doc, "This means district TAY context is useful for district comparison, but district TAY values should not be summed and treated as the council total.")
    add_bullet(doc, "Official district views exclude non-official labels such as Unassigned when comparing districts.")
    add_table(
        doc,
        ["Program view", "Published eligibility basis"],
        [
            ["Packs", "Grades K-5 or ages 5-10"],
            ["Troops", "Grades 6-12 or ages 11-17"],
            ["Crews", "Grades 9-12 or ages 14-20"],
            ["Ships", "Grades 9-12 or ages 14-20"],
            ["Posts", "Grades 9-12 or ages 14-20"],
        ],
        [2600, 6760],
    )
    add_callout(
        doc,
        "Program TAY is an estimate",
        "School sources do not publish enrollment by individual grade. The dashboard assumes even enrollment across each published span. Program populations overlap; fifth-grade and age-13 exceptions cannot be isolated; Exploring Clubs are not represented in Posts; and most school rows do not cover ages 18-20. Do not add program TAY estimates together.",
    )

    add_heading(doc, "9. Training And Safeguarding Logic", 1)
    add_table(
        doc,
        ["Question", "How the dashboard answers it"],
        [
            ["Who is direct-contact?", "The source workbook's Direct Contact column is converted from YES/NO to a boolean."],
            ["Who is trained?", "The Training tab's Trained column is converted from YES/NO to a boolean."],
            ["Is SYT current?", "SYT applies to all leaders. Missing or expired SYT is flagged for all leader rows, not only direct-contact rows."],
            ["Is Hazardous Weather current?", "For direct-contact leaders, missing or expired Hazardous Weather is flagged."],
            ["Is BALOO needed?", "Direct-contact Pack leaders require BALOO; missing or expired BALOO is flagged."],
            ["Is IOLS missing?", "Direct-contact Troop leaders are flagged when mandatory code S11 remains present."],
            ["What about all-leader views?", "Existing safety dates are shown for all rows. SYT issues apply to all leaders; Hazardous Weather, BALOO, and IOLS issues are evaluated where role and unit type make them applicable."],
        ],
        [2800, 6560],
    )
    add_body(
        doc,
        "Training code names are looked up from the Training Codes tab so the SYT page can translate course codes into readable course names where possible.",
    )

    add_heading(doc, "Camping readiness", 2)
    add_body(
        doc,
        "The Pack and Troop Camping Readiness pages group published Training-tab rows by District and unit. They are roster-readiness and training follow-up aids, not campout approvals, attendance rosters, or replacements for council and Guide to Safe Scouting requirements.",
    )
    add_table(
        doc,
        ["Readiness signal", "Current dashboard rule"],
        [
            ["Pack BALOO coverage", "At least one registered Pack leader has BALOO recorded as Yes or as a recognizable completion date."],
            ["Pack Hazardous Weather coverage", "At least one direct-contact Pack leader has a recognizable Hazardous Weather date on or after the published data generated date."],
            ["Troop IOLS coverage", "At least one direct-contact Troop leader does not have mandatory code S11 outstanding. IOLS is inferred because the published training row has no separate IOLS completion field."],
            ["Troop Hazardous Weather coverage", "At least one direct-contact Troop leader has a recognizable Hazardous Weather date on or after the published data generated date."],
            ["Overnight gap", "Either required roster-level signal is absent. Both signals must be present for the page to show Ready based on roster."],
        ],
        [3000, 6360],
    )
    add_callout(
        doc,
        "Confirm the attending leaders",
        "BALOO, IOLS, and Hazardous Weather status is evaluated from the published roster snapshot. Hazardous Weather is available only for direct-contact rows, and the pages do not know which leaders will attend a particular outing. Confirm the actual attending leadership before treating a flag as final.",
    )

    add_heading(doc, "10. Membership Intelligence Signals", 1)
    add_body(
        doc,
        "Membership Intelligence is a dashboard-created prioritization view. It combines source workbook membership and unit-health metrics with monday.com school, prospect, and renewal context.",
    )
    add_table(
        doc,
        ["Signal input", "How it is counted"],
        [
            ["Schools", "School rows whose Scouting District includes the district."],
            ["TAY", "Council view sums attributed school TAY; program views sum the estimated grade/age-eligible portion."],
            ["Schools without unit", "Council view uses blank Unit Associated; program views require an eligible school without an associated unit matching the selected program."],
            ["Hot prospects", "Prospect rows whose District includes the district."],
            ["Stuck prospects", "Prospect rows where Step 1 status is Stuck."],
            ["Renewal follow-up", "Renewal rows where Posted is not Completed."],
        ],
        [2600, 6760],
    )
    add_callout(
        doc,
        "Important interpretation",
        "The Membership Intelligence priority score is not an official workbook metric. It is a dashboard sorting aid that highlights districts where low TAY penetration, membership decline, unit health risk, stuck prospects, and renewal follow-up overlap.",
    )

    add_heading(doc, "11. Commissioner Coverage", 1)
    add_body(
        doc,
        "Commissioner coverage comes from the Commissioners tab and assignment-related tabs in the Council dashboard workbook.",
    )
    add_table(
        doc,
            ["Coverage value", "Calculation"],
            [
            ["Registered commissioners", "Unique commissioner names from the Commissioners tab after whitespace and case normalization."],
            ["Workbook commissioner records", "Raw rows in the Commissioners tab before deduplication."],
            ["Duplicate commissioner records", "Workbook commissioner records minus unique commissioner names."],
            ["Unit commissioners", "Unique people with at least one Unit Commissioner role, counted once even if the person appears multiple times."],
            ["Commissioners trained", "Unique commissioners marked trained divided by unique commissioners where training status is known."],
            ["With assignments", "Unique commissioners with assigned units divided by unique commissioners."],
            ["Assigned units", "Rows in the Assigned tab where Assigned is yes."],
        ],
        [2600, 6760],
    )

    add_heading(doc, "12. Practical Caveats", 1)
    for text in [
        "The published dashboard is only as current as the most recent successful local refresh and GitHub Pages deployment.",
        "The ? buttons provide quick panel context in the website. They are not a replacement for this guide or the Markdown data dictionary.",
        "Workbook labels and sheet names matter. If a source workbook changes structure, the refresh may need a code update.",
        "Expiration checks on the Training and SYT pages use the viewer browser's current date.",
        "Freshness timestamps are displayed in the viewer browser's local timezone, so the same data snapshot can show different clock times to viewers in different timezones.",
        "The SYT detail page flags SYT for all leaders. It flags Hazardous Weather, BALOO, and IOLS only when they are required by the leader's role and unit type.",
        "The Pack and Troop Camping Readiness pages evaluate roster-level training signals from the published snapshot; they do not confirm campout attendance or approval.",
        "monday.com district labels can include operational labels that are not official dashboard districts; official district charts filter those out.",
        "Service Area filters and district leadership are based on the run-specific monday.com Service Areas capture, not on source workbook columns.",
        "Program TAY values are dashboard estimates based on published school grade or age spans and must not be added together.",
        "Popcorn participation counts every published unit row in the denominator and only rows marked Committed in the numerator.",
        "Popcorn contact names, email addresses, and phone numbers are excluded from the public JSON and dashboard.",
        "The Markdown data dictionary remains the best place for exact formulas and implementation details.",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "Appendix: Quick Formula Reference", 1)
    add_table(
        doc,
        ["Metric", "Formula"],
        [
            ["YoY %", "YoY delta divided by same-month-last-year membership."],
            ["At-risk rate", "Units with Unit Metric 0-2 divided by total units."],
            ["Healthy rate", "Units with Unit Metric 4-5 divided by total units."],
            ["Assigned %", "Assigned units divided by total units."],
            ["Council training %", "District training rates weighted by district units."],
            ["Volunteers", "Count of unique nonblank Training roster Member IDs, published by the source workbook on the Membership tab."],
            ["Council SYT %", "District SYT rates weighted by district members."],
            ["Council Youth / TAY", "Council youth divided by raw school-row TAY total."],
            ["District Youth / TAY", "District youth divided by school TAY attributed to that district."],
            ["Program Youth / TAY", "Actual youth from units of the selected program divided by estimated grade/age-eligible school TAY."],
            ["Retention", "(Current members - members new in the prior 12 months) divided by same-month prior-year members; results above 100% are valid, and displayed values round to the nearest whole percent."],
            ["Pack camping readiness", "At least one Pack leader with BALOO recorded and at least one direct-contact Pack leader with current Hazardous Weather."],
            ["Troop camping readiness", "At least one direct-contact Troop leader without S11 outstanding and at least one direct-contact Troop leader with current Hazardous Weather."],
            ["Popcorn Participation", "Committed Popcorn unit rows divided by all Popcorn unit rows."],
            ["Registered commissioners", "Unique normalized commissioner names."],
            ["Unit commissioners", "Unique normalized people with at least one Unit Commissioner role."],
            ["Service Area", "Run-specific monday.com Service Areas hierarchy applied after official district normalization."],
        ],
        [2700, 6660],
    )

    add_heading(doc, "Where to look for exact source details", 2)
    add_reference_bullet(doc, "`DASHBOARD_DATA_DICTIONARY.md` documents the full data dictionary and formulas.")
    add_reference_bullet(doc, "`IMPLEMENTATION_RUNBOOK.md` documents technical setup, source acquisition, automation, validation, and rebuild requirements.")
    add_reference_bullet(doc, "`update_daily.zsh` documents the automation order.")
    add_reference_bullet(doc, "`panel-help.js` documents the website's active ? help popover behavior.")
    add_reference_bullet(doc, "`refresh_monday_data.py` documents the monday.com workbook/API extraction.")
    add_reference_bullet(doc, "`tools/build_unit_level_dashboard.py` documents the Unit-Level Detail data build.")
    add_reference_bullet(doc, "`tools/build_fall_recruitment_dashboard.py` documents the Cub Scout JSN data build.")
    add_reference_bullet(doc, "`work/renewal_recreation/build_renewal_board_data.py` documents the Renewal Status data build.")
    add_reference_bullet(doc, "`work/commissioner_site/build_site.py` documents the source workbook parsing and `latest.json` formulas.")

    OUT_DIR.mkdir(exist_ok=True)
    doc.save(OUT_FILE)
    return OUT_FILE


if __name__ == "__main__":
    print(build_doc())
